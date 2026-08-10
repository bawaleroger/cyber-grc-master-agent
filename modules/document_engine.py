
import re
from pypdf import PdfReader

NORMES_KEYWORDS = {
    "ISO 27001:2022": ["27001","isms","annexe a","soa","smi"],
    "ISO 42001:2023": ["42001","ia responsable","intelligence artificielle","llm","modele ia","ai act"],
    "NIS2": ["nis2","directive 2022/2555"],
    "DORA": ["dora","2022/2554","rts","ict"],
    "RGPD": ["rgpd","gdpr","dpi a","registre traitement","art.30","art.35"],
    "PCI DSS 4.0.1": ["pci","dss","cde","carte bancaire"],
    "BCEAO/UEMOA": ["bceao","uemoa","cobac","beac","gim","art.34","007-09-2017","lcb/ft"],
    "NIST CSF 2.0": ["nist","csf","pr.ac","sp 800-53"],
    "KIT DE CADRAGE": ["kit de cadrage","pv de cadrage","raci","gantt","lettre officielle","liste des documents"]
}

def extract_text(file):
    text = ""
    try:
        if file.name.endswith(".pdf"):
            reader = PdfReader(file)
            for p in reader.pages:
                text += (p.extract_text() or "") + "\n"
        elif file.name.endswith(".docx"):
            import docx
            doc = docx.Document(file)
            # Para + tables
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
        else:
            text = file.getvalue().decode("utf-8", errors="ignore")
    except Exception as e:
        text = f"[ERREUR EXTRACTION {file.name}: {e}]"
    # Clean excessive newlines for preview but keep for parsing
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text

def classify_document(text):
    tl = text.lower()
    scores = {}
    for norme, kws in NORMES_KEYWORDS.items():
        scores[norme] = sum(1 for kw in kws if kw.lower() in tl)
    best = max(scores, key=scores.get) if scores else "GENERAL GRC"
    if scores.get(best,0) == 0:
        best = "GENERAL GRC"
    # Type detection
    if any(x in tl for x in ["termes de reference","t.d.r","dao","cahier des charges","appel d'offres"]):
        dtype = "TDR/DAO"
    elif "kit de cadrage" in tl or "taches a faire" in tl or "liste des documents" in tl:
        dtype = "KIT CADRAGE / CHECK-LIST"
    elif any(x in tl for x in ["p.s.s.i","politique de securite","procedure","registre","soa"]):
        dtype = "LIVRABLE / REFERENTIEL"
    else:
        dtype = "PREUVE / DOC GEN"
    return dtype, best, scores

def parse_tdr(text):
    # Nettoyage intelligent
    clean = text.strip()
    # Extraction objectif - cherche KIT DE CADRAGE ou OBJECTIF ou contexte
    objectif_match = re.search(r'(KIT DE CADRAGE.*?(?:\n\n|\Z))', clean, re.I | re.S)
    if objectif_match:
        objectif_raw = objectif_match.group(1)[:1500]
    else:
        # fallback premiers 800 chars utiles
        objectif_raw = re.sub(r'\s+', ' ', clean[:1000])

    # Perimetre - detecte BCEAO etc
    perimetre = "SI complet + Cloud M365/AWS + Infra reseau + AD + Applications metier + Systemes IA + Prestataires critiques + OT si applicable"
    if "bceao" in clean.lower() or "bank" in clean.lower() or "uemoa" in clean.lower():
        perimetre += " | Perimetre BCEAO/UEMOA: SIB, Monétique GIM-UEMOA, SWIFT, e-Banking, LCB/FT"

    # Secteur
    secteur = "Finance / Banque UEMOA" if "bceao" in clean.lower() or "bank" in clean.lower() else "Multi-secteur critique"

    # Normes detectees
    normes = []
    for norme, kws in NORMES_KEYWORDS.items():
        if norme == "KIT DE CADRAGE": continue
        if any(kw.lower() in clean.lower() for kw in kws):
            normes.append(norme)
    if not normes:
        normes = ["ISO 27001:2022","ISO 42001:2023","NIS2","DORA","BCEAO/UEMOA","PCI DSS 4.0.1"]

    # Livrables extraits
    livrables = []
    if "kit de cadrage" in clean.lower():
        # Parse liste à partir du doc
        for line in clean.split("\n"):
            line=line.strip()
            if len(line)>5 and len(line)<120:
                if any(k in line.lower() for k in ["rapport","pv","lettre","raci","gantt","liste","matrice","questionnaire"]):
                    livrables.append(line)
    if not livrables:
        livrables = ["Rapport de Cadrage","PV de Reunion de Cadrage","Lettre Officielle + Liste Docs a Demander","RACI + Planning Gantt Excel","Questionnaire 27001/42001 350 questions","Matrice Gap Analysis","Registre Risques","Plan d'action chiffre XOF"]

    return {
        "objectif": objectif_raw,
        "perimetre": perimetre,
        "secteur": secteur,
        "normes": normes,
        "livrables": livrables,
        "contraintes": "Delais regulateur BCEAO + Confidentialite + Disponibilite equipes metiers + Acces docs probants",
        "raw_preview": clean[:2000]
    }
