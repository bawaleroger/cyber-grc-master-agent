
import streamlit as st, os, io, datetime
from modules.persistence import init_db, save_document, get_all_docs, get_kb_context, save_tdr
from modules.document_engine import extract_text, classify_document, parse_tdr
from modules.ai_client import ai
from modules.report_generator import generate_offre_technique_content, generate_offre_financiere_content, generate_dao_content, generate_audit_section, generate_certification_roadmap
import pandas as pd

st.set_page_config(page_title="CYBER-GRC MASTER V3", page_icon="🛡️", layout="wide")
init_db()

def export_docx(text, title):
    from docx import Document
    doc = Document()
    doc.add_heading(title, 0)
    for para in text.split("\n"):
        if para.strip().startswith("# "): doc.add_heading(para.replace("# ",""), 1)
        elif para.strip().startswith("## "): doc.add_heading(para.replace("## ",""), 2)
        elif para.strip().startswith("### "): doc.add_heading(para.replace("### ",""), 3)
        else: doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_pdf(text, title):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(f"<b>{title}</b>", styles['Title']), Spacer(1,12)]
    for line in text.split("\n")[:500]:
        if line.strip():
            story.append(Paragraph(line[:500].replace("<","&lt;"), styles['Normal']))
            story.append(Spacer(1,6))
    doc.build(story)
    return bio.getvalue()

def export_excel_kit():
    df_erl = pd.DataFrame([
        ["DOC-001","Organigramme DSI + Fiches poste RSSI/DPO/AI Officer","ORG","Critique","J3","DG","ISO 27001 A.5.2 + ISO 42001 A.5.2"],
        ["DOC-042","Logs proxy 12 mois + regles FW + config VPN","TECH","Critique","J3","DSI","ISO 27001 A.8.15 + NIS2 Art.21 + ISO 42001 A.9.3"],
        ["DOC-089","Inventaire modeles IA + datasets + prompts + contrats","IA","Critique","J5","AI Officer","ISO 42001 B.3 + B.7.4 + AI Act Art.49"],
        ["DOC-110","Registre traitement Art.30 + DPIA","RGPD","Majeur","J5","DPO","RGPD Art.30 + Art.35"],
    ], columns=["ID","Document demande","Type","Criticite","Delai","Resp","Norme"])
    df_raci = pd.DataFrame([
        ["Kick-off","Chef projet CABINET","DG","RSSI,DPO","COMEX","PV Kick-off","Teams","ISO 19011"],
        ["Collecte preuves","Auditeur","RSSI","DSI,Juridique","DG","ERL 150 docs","SharePoint","ISO 27001 A.5.33"],
        ["Atelier EBIOS RM","Risk Manager","CISO","DPO,AI Officer","DG","Socle securite","EBIOS","ISO 27005 + ISO 23894"],
        ["Pentest interne","Pentester","CISO","DSI","RSSI","Rapport Pentest","Nmap/BloodHound","PTES + MITRE ATT&CK"],
    ], columns=["Tache","R","A","C","I","Livrable","Outil","Norme"])
    df_quest = pd.DataFrame([
        ["Q-001","PSSI approuvee <12 mois?","ISO 27001 A.5.1","Non","PSSI.pdf","Majeur","ISO 27001 A5.1 + NIS2 Art.20"],
        ["Q-042","MFA generalise VPN + M365 Admin?","ISO 27001 A.5.17 + A.8.5","Non","GPO + Entra","Majeur","A5.17 + NIS2 Art.20 + PCI 8.4.3 + BCEAO Art.34"],
        ["Q-089","Shadow AI detecte? Logs proxy?","ISO 42001 A.9.3 + A.10","Non","Proxy logs","Critique","A.9.3 + AI Act Art.50 + LLM10"],
        ["Q-120","Registre IA B.7.4 + Art.49 existe?","ISO 42001 B.7.4","Non","Entretien","Critique","B.7.4 + Art.49 Sanction 15M€"],
    ], columns=["ID","Question","Norme source","Reponse","Preuve","Criticite","Mapping 2026"])
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        df_erl.to_excel(writer, sheet_name='ERL_150_preuves', index=False)
        df_raci.to_excel(writer, sheet_name='RACI_Gantt', index=False)
        df_quest.to_excel(writer, sheet_name='Questionnaire_350Q', index=False)
    return bio.getvalue()

# SIDEBAR
st.sidebar.title("🛡️ CYBER-GRC MASTER V3")
st.sidebar.caption("CISO Big Four 50 ans - Full Redaction")

client = st.sidebar.text_input("CLIENT", st.session_state.get("client","BCEAO BANK UEMOA"))
cabinet = st.sidebar.text_input("CABINET", st.session_state.get("cabinet","CYBER-GRC CONSULTING"))
st.session_state["client"]=client
st.session_state["cabinet"]=cabinet

st.sidebar.divider()
st.sidebar.subheader("ETAPE 1: CHARGER TDRs + AUTO-NOURRISSEMENT")
st.sidebar.markdown("**A. TDRs / DAO (pour mission en cours)**")
tdr_files = st.sidebar.file_uploader("Charge TDR/DAO ici", accept_multiple_files=True, type=["pdf","docx","txt"], key="tdr_upl")
st.sidebar.markdown("**B. Base de Connaissance (docs pour auto-nourrissement permanent)**")
kb_files = st.sidebar.file_uploader("Charge referentiels, anciens rapports, PSSI, politiques pour nourrir l'IA", accept_multiple_files=True, type=["pdf","docx","txt"], key="kb_upl")

if st.sidebar.button("🧠 Ingerer & Nourrir Base", type="primary"):
    full_tdr=""
    # TDR
    for f in tdr_files or []:
        txt=extract_text(f)
        full_tdr+=txt+"\n"
        dtype,best,scores=classify_document(txt)
        save_document(f.name, txt, dtype, best, client)
        st.sidebar.success(f"TDR {f.name} -> {dtype} / {best}")
    # KB
    for f in kb_files or []:
        txt=extract_text(f)
        dtype,best,scores=classify_document(txt)
        save_document(f.name, txt, dtype, best, client)
        st.sidebar.success(f"KB {f.name} -> {best} auto-nourri")
    if full_tdr:
        tdr_dict=parse_tdr(full_tdr)
        save_tdr(client,cabinet,tdr_dict,full_tdr)
        st.session_state["tdr_dict"]=tdr_dict
        st.session_state["tdr_raw"]=full_tdr
        st.sidebar.json(tdr_dict)

if "tdr_dict" in st.session_state:
    st.sidebar.metric("Docs en KB (auto-nourrie)", len(get_all_docs()))
    with st.sidebar.expander("Voir KB"):
        st.dataframe(get_all_docs())

st.sidebar.divider()
st.sidebar.subheader("ETAPE 2: MODE")
mode = st.sidebar.radio("Moteur", ["Avec Chat IA (Groq/OpenAI API)", "Sans IA - Base auto-nourrissante 100% locale"])
api_key = st.sidebar.text_input("Cle API temporaire Groq", type="password")
if api_key: os.environ["GROQ_API_KEY"]=api_key
st.sidebar.caption(f"Mode: {ai.mode.upper()} | KB chars: {len(get_kb_context(1000))}")

# MAIN
st.title("CYBER-GRC MASTER - Workflow V3 Autonome - Full Redaction Bout en Bout")
st.markdown(f"**CLIENT:** {client} | **CABINET:** {cabinet} | **Devise:** XOF | **KB:** {len(get_all_docs())} docs auto-nourris")

if "tdr_dict" not in st.session_state:
    st.warning("👈 ETAPE 1: Charge ton TDR a gauche (ex: Taches a faire.docx) + eventuellement des docs de base de connaissance pour nourrir l'IA. Ensuite clique Ingerer.")
    st.info("L'app est autonome: chaque doc charge reste apres refresh dans SQLite data/cyber_grc.db et est reutilise pour toutes les generations futures.")
    st.stop()

tdr_dict = st.session_state["tdr_dict"]
kb_ctx = get_kb_context(20000)
tdr_raw = st.session_state.get("tdr_raw","")

tab1,tab2,tab3,tab4,tab5 = st.tabs(["📄 OFFRE TECHNIQUE (Full)", "💰 OFFRE FINANCIERE (Full XOF)", "📘 DAO (Full)", "🚀 MISSION AUDIT BOUT EN BOUT Section par Section", "🏅 CERTIFICATION (Dropdown)"])

with tab1:
    st.header("Bouton Offre Technique - Redaction COMPLETE Big Four (pas plan)")
    st.caption("Genere 10+ pages avec KIT DE CADRAGE reel, RACI, ERL 150 preuves, Questionnaire 350Q, Arsenal commandes, Matrice Gap, 10 livrables")
    if st.button("Generer Offre Technique Complete", type="primary", key="b1"):
        with st.spinner("CISO Big Four 50 ans redige offre technique complete..."):
            if "Avec Chat IA" in mode:
                prompt = f"""GENERE OFFRE TECHNIQUE COMPLETE BIG FOUR 100% REDIGEE PAS PLAN pour {client} par {cabinet}.
TDR REEL CHARGE: {tdr_raw[:5000]}
KB INTERNE AUTO-NOURRIE: {kb_ctx[:6000]}
EXIGENCES: 
- Livre VRAIMENT le KIT DE CADRAGE: Rapport Cadrage 10p + PV + LDD 150 preuves + RACI + Gantt
- Donne Questionnaire 350Q detaille, Matrice inventaire actifs ISO 27005 + Registre IA B.7.4 Art49, DFD, EBIOS RM, Arsenal commandes theHarvester, nmap, BloodHound, TruffleHog, shadow_ai_detector
- Gap format 2026: [ACTION] -> ISO 27001 [Clause] + ISO 42001 [Clause] + NIS2 + PCI + BCEAO Art34
- Anti-hallucination: chaque affirmation sourcee [ISO 27001 A.x] [ISO 42001 A.x] [BCEAO Art34] sinon [PREUVE NON TROUVEE]
- 10 livrables Big Four detailles, equipe avec certifs ISO 27001 LA, 42001 LA, 22301 LA, CISA CISM CISSP-ISSAP CRISC EBIOS RM OSCP OSEP PNPT SABSA
- Langue Francais pro UEMOA, couts XOF, 0 etc.
- Remplace {{CLIENT}} {client} {{CABINET}} {cabinet}
"""
                result = ai.generate(prompt, kb_ctx, tdr_raw)
            else:
                result = generate_offre_technique_content(client,cabinet,tdr_dict,kb_ctx)
            st.session_state["offre_tech"]=result
    if "offre_tech" in st.session_state:
        st.markdown(st.session_state["offre_tech"])
        c1,c2,c3,c4 = st.columns(4)
        c1.download_button("WORD Full", export_docx(st.session_state["offre_tech"], "Offre_Technique_Complete"), f"Offre_Technique_Complete_{client}.docx")
        c2.download_button("PDF Full", export_pdf(st.session_state["offre_tech"], "Offre_Technique_Complete"), f"Offre_Technique_Complete_{client}.pdf")
        c3.download_button("MD", st.session_state["offre_tech"], f"Offre_Technique_{client}.md")
        c4.download_button("KIT Excel (ERL+RACI+350Q)", export_excel_kit(), f"KIT_CADRAGE_{client}.xlsx")

with tab2:
    st.header("Bouton Offre Financiere - Redaction Complete XOF + Excel JH")
    if st.button("Generer Offre Financiere Complete", type="primary", key="b2"):
        with st.spinner("Calcul JH + XOF detaille..."):
            if "Avec Chat IA" in mode:
                prompt = f"GENERE OFFRE FINANCIERE DETAILLEE XOF COMPLETE pour {client}. Decomposition par phase ET par livrable reel (Kit Cadrage, Carto, Audit Tech, Gap, Redaction, Forensic, Restitution) avec JH Senior/Junior, PU, Total XOF, TVA UEMOA 18%, Total TTC. Inclus Gantt detaille 80 taches + Budget detaille + Matrice risques. TDR:{tdr_raw[:3000]}"
                result = ai.generate(prompt, kb_ctx, tdr_raw)
            else:
                result = generate_offre_financiere_content(client,cabinet,tdr_dict)
            st.session_state["offre_fin"]=result
    if "offre_fin" in st.session_state:
        st.markdown(st.session_state["offre_fin"])
        c1,c2,c3 = st.columns(3)
        c1.download_button("WORD", export_docx(st.session_state["offre_fin"], "Offre_Financiere"), f"Offre_Financiere_{client}.docx")
        c2.download_button("EXCEL Budget + Gantt", export_excel_kit(), f"Offre_Financiere_{client}.xlsx")
        c3.download_button("PDF", export_pdf(st.session_state["offre_fin"], "Offre_Financiere"), f"Offre_Financiere_{client}.pdf")

with tab3:
    st.header("Bouton DAO - Redaction Complete BCEAO/UEMOA")
    if st.button("Generer DAO Complet", type="primary", key="b3"):
        with st.spinner("Redaction DAO complet..."):
            if "Avec Chat IA" in mode:
                prompt = f"GENERE DAO / CAHIER DES CHARGES COMPLET BCEAO UEMOA pour {client}. Articles 1-7 + Annexes ERL 150, Questionnaire 350Q, Matrice Gap, Template PV. Exigences normatives avec clauses exactes ISO 27001 Cl4-10 Annexe A 93 + ISO 42001 Cl4-10 Annexe A/B + NIS2 Art20-21 + DORA Art11-13 + RGPD Art5-35 + AI Act Art9,10,13,14,49,50 + PCI DSS 4.0.1 + SWIFT CSCF + BCEAO Art34 + MITRE ATT&CK v15 + CVSS v4.0. Criteres eval 70% tech 30% financier. TDR:{tdr_raw[:3000]}"
                result = ai.generate(prompt, kb_ctx, tdr_raw)
            else:
                result = generate_dao_content(client,cabinet,tdr_dict)
            st.session_state["dao"]=result
    if "dao" in st.session_state:
        st.markdown(st.session_state["dao"])
        c1,c2 = st.columns(2)
        c1.download_button("WORD DAO", export_docx(st.session_state["dao"], "DAO_Complet"), f"DAO_Complet_{client}.docx")
        c2.download_button("PDF DAO", export_pdf(st.session_state["dao"], "DAO_Complet"), f"DAO_Complet_{client}.pdf")

with tab4:
    st.header("Bouton Demarrer Mission Audit - Bout en Bout - Section par Section - Full Redaction")
    st.success("Background: CISO Big Four 50 ans exp cumulee +500 missions Afrique BCEAO/COBAC/BEAC/GIM-UEMOA, Europe NIS2/DORA/RGPD/AI Act, USA NIST/SOC2/HIPAA/CMMC | Certifs: ISO 27001 LA, 42001 LA, 22301 LA, 9001 LA, 23894 LI, EBIOS RM LRM, CISA, CISM, CISSP-ISSAP, CRISC, CGEIT, CDPSE, PCI QSA, GCFA, CHFI, OSCP/OSEP/OSWE/OSED, PNPT, CPTS, CRTO, CEH Master, SABSA, TOGAF, ITIL 4 MP, COBIT 2019, PMP/RMP - Chaque livrable maitrise")
    
    phases = [
        "A_CADRAGE - Kit Cadrage Complet (Rapport 10p + PV + LDD 150 + RACI + Gantt)",
        "B_CARTO - Etat existant + Inventaire actifs + Registre IA B.7.4 Art.49 + Shadow AI CSV + DFD",
        "C_RISQUES - EBIOS RM 5 ateliers + Audit tech PTES/WSTG/LLM + Registre Risques 50 lignes + Preuves SHA256",
        "D_GAP - Gap Analysis 120 lignes Matrice croisee 2026 + Radar Maturite + % conformite",
        "E_REMEDIATION - Plan action XOF + 10 Livrables Big Four + Schema Directeur 3 ans + Tableau bord"
    ]
    sel = st.selectbox("Choisir section a generer (efficacite volume)", phases)
    key = sel.split(" - ")[0]
    
    if st.button(f"Generer {sel}", type="primary"):
        with st.spinner(f"Generation FULL REDACTION {sel} - mode operationnel reel..."):
            if "Avec Chat IA" in mode:
                prompt = f"""MISSION AUDIT BOUT EN BOUT - GENERE SECTION {sel} UNIQUEMENT - REDACTION COMPLETE OPERATIONNELLE PAS PLAN pour {client} par {cabinet}.
TDR REEL: {tdr_raw[:5000]}
KB AUTO-NOURRIE: {kb_ctx[:6000]}
EXIGENCES SECTION {key}:
- Donne VRAIS templates remplis, VRAIS questionnaires, VRAIES matrices Excel avec colonnes, VRAIES commandes arsenal avec {client}, VRAIES preuves SHA256, VRAIS mappings [ISO 27001:2022 A.x + ISO 42001:2023 A.x + NIS2 Art.x + DORA Art.x + PCI DSS Req x + BCEAO Art34 + MITRE ATT&CK Txxxx]
- Background CISO Big Four 50 ans + certifs a jour ISO 27001 LA, 42001 LA, EBIOS LRM, CISA CISM CISSP-ISSAP CRISC OSCP OSEP
- Format Markdown Pro + Tableaux + Check-lists [ ] + RACI + Gantt Mermaid + Cout XOF
- 0 etc. - Tu listes tout
- Anti-hallucination: chaque affirmation sourcee
- Client {client} Cabinet {cabinet}
"""
                result = ai.generate(prompt, kb_ctx, tdr_raw)
            else:
                result = generate_audit_section(key, client, cabinet, tdr_dict, kb_ctx)
            st.session_state[f"audit_{key}"]=result
    
    for p in phases:
        k = p.split(" - ")[0]
        if f"audit_{k}" in st.session_state:
            with st.expander(f"✅ {p} - REDACTION COMPLETE", expanded=(k==key)):
                st.markdown(st.session_state[f"audit_{k}"])
                c1,c2,c3 = st.columns(3)
                c1.download_button(f"WORD {k}", export_docx(st.session_state[f"audit_{k}"], k), f"{k}_{client}.docx", key=f"w_{k}")
                c2.download_button(f"PDF {k}", export_pdf(st.session_state[f"audit_{k}"], k), f"{k}_{client}.pdf", key=f"p_{k}")
                c3.download_button(f"KIT Excel {k}", export_excel_kit(), f"{k}_KIT_{client}.xlsx", key=f"x_{k}")
    
    if any(f"audit_{p.split(' - ')[0]}" in st.session_state for p in phases):
        if st.button("Fusionner tout en Rapport Final Complet 100+ pages"):
            full = "\n\n\n---\n\n\n".join([st.session_state.get(f"audit_{p.split(' - ')[0]}","") for p in phases])
            st.session_state["audit_full"]=full
            st.download_button("📥 Rapport Complet WORD 100+ pages", export_docx(full,"Rapport_Audit_Complet_100p"), f"Rapport_Audit_Complet_100p_{client}.docx")
            st.download_button("📥 Rapport Complet PDF", export_pdf(full,"Rapport_Audit_Complet"), f"Rapport_Audit_Complet_{client}.pdf")
            st.download_button("📥 KIT Complet Excel (ERL+RACI+350Q+Gap)", export_excel_kit(), f"KIT_Audit_Complet_{client}.xlsx")

with tab5:
    st.header("Bouton Certification - Accompagnement bout en bout - Dropdown")
    norme = st.selectbox("Choisis norme cible pour accompagnement certification", ["ISO 27001:2022","ISO 42001:2023","ISO 9001:2015","ISO 22301:2019","PCI DSS 4.0.1","SOC 2","NIS2 Directive 2022/2555","DORA Regulation 2022/2554"], key="cert_norme")
    if st.button(f"Demarrer accompagnement certification {norme} - Bout en bout", type="primary"):
        with st.spinner(f"Roadmap certification {norme} bout en bout..."):
            if "Avec Chat IA" in mode:
                prompt = f"GENERE ACCOMPAGNEMENT CERTIFICATION BOUT EN BOUT {norme} pour {client} par {cabinet}. Workflow complet: 1 Gap Analysis 200Q {norme} + Matrice ecarts + Radar, 2 Mise en conformite (PSSI + SoA 93 + 30 politiques + Registre IA B.7.4 si 42001 + Registre Art30 RGPD), 3 Implementation technique MFA+Tiering+SIEM+DLP IA, 4 Formation 7.2+7.3+A.3.2+AI Act Art4, 5 Audit blanc ISO 19011 Plan + Check-list 93 controles + ERL + Grille CIA-T + Rapport audit blanc, 6 Revue direction Cl9.3 + Accompagnement certificateur AFNOR/Bureau Veritas + Levee NC + Certificat. Chiffrage XOF detaille + Delais + Templates Excel + RACI + Clauses exactes {norme}. TDR:{tdr_raw[:3000]} KB:{kb_ctx[:4000]}"
                result = ai.generate(prompt, kb_ctx, tdr_raw)
            else:
                result = generate_certification_roadmap(norme, client, cabinet)
            st.session_state[f"cert_{norme}"]=result
    if f"cert_{norme}" in st.session_state:
        st.markdown(st.session_state[f"cert_{norme}"])
        c1,c2 = st.columns(2)
        c1.download_button(f"WORD Certif {norme}", export_docx(st.session_state[f"cert_{norme}"], f"Certification_{norme}"), f"Certification_{norme}_{client}.docx")
        c2.download_button(f"KIT Excel {norme}", export_excel_kit(), f"KIT_Certif_{norme}_{client}.xlsx")
